from docx import Document
import pandas as pd
import math
import numpy as np
import os
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE
from docx.dml.color import ColorFormat
from docx.enum.dml import MSO_COLOR_TYPE
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.text import WD_COLOR
from docx.shared import Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_ALIGN_PARAGRAPH

directory = "C:\\Users\\t8455242\\Desktop\\WinPython-64bit-3.6.2.0Qt5\\notebooks\\ניתוח סוציומטרי\\dir"

word_index = {
    'תוך אישי': (3, 4),
    'מדעי אקדמי': (4, 4),
    'מדעי יישומי': (7, 4),
    'תפקוד בחברה': (3, 8),
    'ניהול והתנהלות': (4, 8),
    'בטחוני': (7, 8),
    'הובלה': (3, 14),
    'ערכי': (4, 14),
    'אחריות': (5, 11),
    'מצוינות': (6, 11),
    'יושרה': (7, 11),
    'העזה': (5, 14),
    'שליחות': (6, 14),
    'דרך ארץ': (7, 14),
    'שימור': (9, 14),
    'שיפור': (9, 22),
    'ציטוטים': (9, 44),
    'פרשנות': (9, 74),
    'יעדים': (9, 89),
}

excel_index = {
    'תוך אישי': (1),
    'מדעי אקדמי': (5),
    'מדעי יישומי': (6),
    'תפקוד בחברה': (2),
    'ניהול והתנהלות': (4),
    'בטחוני': (7),
    'הובלה': (3),
    'אחריות': (8),
    'מצוינות': (9),
    'יושרה': (10),
    'העזה': (11),
    'שליחות': (12),
    'דרך ארץ': (13),
}

values = ['אחריות', 'מצוינות', 'יושרה', 'העזה', 'שליחות', 'דרך ארץ']


def num_to_color(name, number):
    color6 = parse_xml(r'<w:shd {} w:fill="00B005"/>'.format(nsdecls('w')))
    color5 = parse_xml(r'<w:shd {} w:fill="94D005"/>'.format(nsdecls('w')))
    color4 = parse_xml(r'<w:shd {} w:fill="FFFF00"/>'.format(nsdecls('w')))
    color3 = parse_xml(r'<w:shd {} w:fill="FFC000"/>'.format(nsdecls('w')))
    color2 = parse_xml(r'<w:shd {} w:fill="FF0000"/>'.format(nsdecls('w')))
    color1 = parse_xml(r'<w:shd {} w:fill="C00000"/>'.format(nsdecls('w')))
    color_mean = parse_xml(r'<w:shd {} w:fill="E7E6E6"/>'.format(nsdecls('w')))
    color = [color_mean, color1, color2, color3, color4, color5, color6]
    if name in values:
        if number == 1:
            return color[1]
        if number == 2:
            return color[0]
        if number == 3:
            return color[6]
    else:
        return color[number]
    print('error')


def create_docx_dir(input_dir, output_dir):
    for filename in os.listdir(input_dir):
        print(filename)
        if filename.endswith(".xlsx"):
            sheet = pd.read_excel(input_dir + "\\" + filename)
            sheet = np.array(sheet)
            negative_points = sheet[:-2, -1]
            positive_points = sheet[:-2, -2]
            document = Document('./Sociometry/Templates/sozio_text.docx')
            table = document.tables[0]
            fill_cells = list(excel_index.keys())
            for value in fill_cells:
                cell = table.cell(int(word_index[value][0]), int(word_index[value][1]))

                # print(excel_index[2])
                index = int(excel_index[value])
                print(sheet[-1, index + 1])
                cell.text = str(int(sheet[-1, index + 1]))
                cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER
                cell._tc.get_or_add_tcPr().append(num_to_color(value, int(sheet[-1, int(excel_index[value]) + 1])))

            text_neg = ""
            text_pos = ""
            for point in negative_points:
                if not type(point) == str:
                    continue
                text_neg = text_neg + "\"" + point + "\"" + "\n\n"
            for point in positive_points:
                if not type(point) == str:
                    continue
                text_pos = text_pos + "\"" + point + "\"" + "\n\n"
            cell = table.cell(int(word_index['שיפור'][0]), int(word_index['שיפור'][1]))
            cell.text = text_neg
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            cell = table.cell(int(word_index['שימור'][0]), int(word_index['שימור'][1]))
            cell.text = text_pos
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            p = document.paragraphs[1]
            name = sheet[0, 0]
            p.text = p.text + " " + str(filename) + " "
            document.save(output_dir + '//' + p.text + '.docx')
